"""
TDD test suite for convert.py — the standalone PDF conversion core.

Covers the three supported conversions (word / excel / ppt) plus CLI/argument
error handling. Tests use a synthetic multi-page PDF built with PyMuPDF so they
run without any external fixture files.
"""
import os
import zipfile

import pytest
import fitz  # PyMuPDF
import openpyxl
from docx import Document
from pptx import Presentation

import convert


@pytest.fixture(scope="module")
def sample_pdf(tmp_path_factory):
    """A 2-page PDF containing plain text on each page."""
    path = tmp_path_factory.mktemp("data") / "sample.pdf"
    doc = fitz.open()
    p1 = doc.new_page()
    p1.insert_text((72, 72), "Hello World Page 1")
    p2 = doc.new_page()
    p2.insert_text((72, 72), "Second Page Content")
    doc.save(str(path))
    doc.close()
    return str(path)


def _docx_text(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# --- Word -------------------------------------------------------------------

def test_convert_word_creates_valid_docx(sample_pdf, tmp_path):
    out = str(tmp_path / "out.docx")
    convert.convert("word", sample_pdf, out)
    assert os.path.exists(out)
    assert zipfile.is_zipfile(out)  # docx is a zip container
    text = _docx_text(out)
    assert "Hello" in text


# --- Excel ------------------------------------------------------------------

def test_convert_excel_creates_valid_xlsx(sample_pdf, tmp_path):
    out = str(tmp_path / "out.xlsx")
    convert.convert("excel", sample_pdf, out)
    assert os.path.exists(out)
    assert zipfile.is_zipfile(out)  # xlsx is a zip container
    wb = openpyxl.load_workbook(out)
    assert len(wb.sheetnames) >= 1


# --- PowerPoint -------------------------------------------------------------

def test_convert_ppt_creates_valid_pptx_with_slide_per_page(sample_pdf, tmp_path):
    out = str(tmp_path / "out.pptx")
    convert.convert("ppt", sample_pdf, out)
    assert os.path.exists(out)
    assert zipfile.is_zipfile(out)  # pptx is a zip container
    prs = Presentation(out)
    assert len(prs.slides) == 2  # one slide per source page


# --- Dispatcher / argument handling ----------------------------------------

def test_convert_unknown_format_raises(sample_pdf, tmp_path):
    out = str(tmp_path / "out.bin")
    with pytest.raises(ValueError):
        convert.convert("csv", sample_pdf, out)


def test_convert_missing_input_raises(tmp_path):
    out = str(tmp_path / "out.docx")
    with pytest.raises(FileNotFoundError):
        convert.convert("word", str(tmp_path / "missing.pdf"), out)


def test_main_wrong_args_returns_nonzero():
    assert convert.main(["word", "only-two-args"]) != 0


def test_main_unknown_format_returns_nonzero(sample_pdf, tmp_path):
    out = str(tmp_path / "out.bin")
    assert convert.main(["csv", sample_pdf, out]) != 0


def test_main_success_returns_zero(sample_pdf, tmp_path):
    out = str(tmp_path / "ok.docx")
    assert convert.main(["word", sample_pdf, out]) == 0
    assert os.path.exists(out)
